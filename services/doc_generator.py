from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
from database.models import User, Vacation

TEMPLATES_DIR = "2025-20260608T221020Z-3-001/2025"
OUTPUT_DIR = "generated_orders"

MONTHS_UA = {
    1: "січня", 2: "лютого", 3: "березня", 4: "квітня",
    5: "травня", 6: "червня", 7: "липня", 8: "серпня",
    9: "вересня", 10: "жовтня", 11: "листопада", 12: "грудня"
}

def ensure_output_dir():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

def format_date_ua(d) -> str:
    """Форматує дату як '15 липня 2025 року'"""
    return f"{d.day} {MONTHS_UA[d.month]} {d.year} року"

def set_font(run, size=12, bold=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=False, spacing_before=0, spacing_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    return p

async def generate_vacation_order(user: User, vacation: Vacation) -> str:
    """Генерує наказ про надання відпустки у форматі школи."""
    ensure_output_dir()
    
    doc = Document()

    # Поля сторінки
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    today = datetime.now()
    today_ua = format_date_ua(today)
    start_ua = format_date_ua(vacation.start_date)
    end_ua = format_date_ua(vacation.end_date)

    # Шапка — назва закладу (замість реальної — заглушка, редагується вручну)
    add_paragraph(doc,
        "КОМУНАЛЬНИЙ ЗАКЛАД ЗАГАЛЬНОЇ СЕРЕДНЬОЇ ОСВІТИ",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
    add_paragraph(doc,
        "(назва закладу)",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

    add_paragraph(doc, spacing_before=6)

    # Номер і дата
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_n = p_num.add_run("НАКАЗ")
    set_font(run_n, size=14, bold=True)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = p_date.add_run(f"від {today_ua}    № ___")
    set_font(run_d, size=12)

    add_paragraph(doc, spacing_before=4)

    # Підстава / вступна частина
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("Про надання щорічної основної відпустки")
    set_font(run_t, size=12, bold=True)

    add_paragraph(doc, spacing_before=6)

    p_base = doc.add_paragraph()
    p_base.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_base.paragraph_format.first_line_indent = Cm(1.25)
    run_b = p_base.add_run(
        f"Відповідно до статті 75 Кодексу законів про працю України, "
        f"на підставі заяви {user.full_name},"
    )
    set_font(run_b, size=12)

    add_paragraph(doc, spacing_before=4)

    # НАКАЗУЮ
    p_nakaz = doc.add_paragraph()
    p_nakaz.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nk = p_nakaz.add_run("НАКАЗУЮ:")
    set_font(run_nk, size=12, bold=True)

    add_paragraph(doc, spacing_before=4)

    # Основний текст
    p_main = doc.add_paragraph()
    p_main.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_main.paragraph_format.first_line_indent = Cm(1.25)
    run_m = p_main.add_run(
        f"1. Надати {user.full_name} щорічну основну оплачувану відпустку "
        f"тривалістю {vacation.days_count} ({"".join(['_'*20])}) календарних днів "
        f"з {start_ua} по {end_ua} включно."
    )
    set_font(run_m, size=12)

    p_control = doc.add_paragraph()
    p_control.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_control.paragraph_format.first_line_indent = Cm(1.25)
    run_c = p_control.add_run(
        "2. Контроль за виконанням даного наказу залишаю за собою."
    )
    set_font(run_c, size=12)

    add_paragraph(doc, spacing_before=12)

    # Підпис директора
    p_dir = doc.add_paragraph()
    p_dir.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_dir = p_dir.add_run("Директор                                    _______________ / ________________")
    set_font(run_dir, size=12)

    add_paragraph(doc, spacing_before=16)

    # Ознайомлення
    p_ack = doc.add_paragraph()
    run_ack = p_ack.add_run("З наказом ознайомлений(а):")
    set_font(run_ack, size=12)

    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_s = p_sign.add_run(
        f"{user.full_name}                        _______________ / ________________\n"
        f"Дата: _______________"
    )
    set_font(run_s, size=12)

    # Зберігаємо
    filename = f"nakaz_vidpustka_{user.full_name.replace(' ', '_')}_{vacation.id}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    
    return filepath


async def generate_order_from_template(
    template_name: str,
    variables: dict,
    output_name: str = None
) -> str:
    """Генерує наказ з шаблону .docx."""
    ensure_output_dir()
    
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Шаблон не знайдено: {template_path}")
    
    doc = Document(template_path)
    
    # Заміна плейсхолдерів у параграфах
    for paragraph in doc.paragraphs:
        for key, value in variables.items():
            placeholder = f"{{{key}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
    
    # Заміна плейсхолдерів у таблицях
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in variables.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
    
    if not output_name:
        output_name = f"nakaz_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    output_path = os.path.join(OUTPUT_DIR, output_name)
    doc.save(output_path)
    
    return output_path


def get_available_templates():
    """Отримати список доступних шаблонів наказів."""
    if not os.path.exists(TEMPLATES_DIR):
        return []
    
    templates = []
    for file in os.listdir(TEMPLATES_DIR):
        if file.endswith('.docx'):
            templates.append(file)
    
    return templates
