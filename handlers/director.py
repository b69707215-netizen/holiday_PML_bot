"""
Хендлери для директора:
- Меню директора (3 кнопки)
- Призначення персоналу по номеру телефону (секретар / заступник)
- Вкладка наказів — схвалені відпустки + генерація docx
- Вкладка відпусток
"""
import os
import re
from datetime import datetime
from aiogram import Router, F
from aiogram.types import Message, CallbackQuery, FSInputFile
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from database import SessionLocal, User, UserRole, Vacation, VacationStatus
from keyboards import (
    director_main_menu, cancel_button,
    confirm_appoint_keyboard, appoint_role_keyboard,
    approved_vacations_keyboard
)
from states import AppointVice
from config import DIRECTOR_ID

router = Router()

def is_director(user_id: int) -> bool:
    return user_id == DIRECTOR_ID

# ─────────────────────────────────────────────────────────
# 1. ПРИЗНАЧИТИ ПЕРСОНАЛ — по номеру телефону
# ─────────────────────────────────────────────────────────

@router.message(F.text == "👤 Призначити персонал")
async def appoint_staff_start(message: Message, state: FSMContext):
    if not is_director(message.from_user.id):
        await message.answer("❌ Тільки для директора.")
        return

    await message.answer(
        "📱 Введіть номер телефону співробітника (у форматі +380...):\n\n"
        "Система знайде людину і запитає яку роль призначити.",
        reply_markup=cancel_button()
    )
    await state.set_state(AppointVice.search_name)
    await state.update_data(appoint_mode="phone")

@router.message(AppointVice.search_name, F.text == "❌ Скасувати")
async def appoint_cancel_text(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Скасовано.", reply_markup=director_main_menu())

@router.message(AppointVice.search_name)
async def appoint_search_by_phone(message: Message, state: FSMContext):
    query = message.text.strip()

    # Нормалізуємо телефон
    phone = re.sub(r"[^\d+]", "", query)
    if not phone.startswith("+"):
        phone = "+" + phone

    async with SessionLocal() as session:
        result = await session.execute(
            select(User).where(User.phone == phone)
        )
        user = result.scalar_one_or_none()

    if not user:
        # Пробуємо по ПІБ як fallback
        async with SessionLocal() as session:
            result = await session.execute(
                select(User).where(User.full_name.ilike(f"%{query}%"))
            )
            users = result.scalars().all()

        if not users:
            await message.answer(
                "❌ Співробітника не знайдено.\n"
                "Перевірте номер або введіть ПІБ:"
            )
            return

        if len(users) == 1:
            user = users[0]
        else:
            lines = [f"• {u.full_name} — {u.phone}" for u in users[:10]]
            await message.answer(
                "Знайдено кілька:\n" + "\n".join(lines) +
                "\n\nУточніть номер телефону:"
            )
            return

    # Знайдено — показуємо і питаємо роль
    role_label = {
        UserRole.DIRECTOR: "Директор",
        UserRole.VICE_PRINCIPAL: "Заступник директора",
        UserRole.SECRETARY: "Секретар",
        UserRole.TEACHER: "Вчитель",
    }.get(user.role, user.role.value)

    await state.update_data(target_id=user.telegram_id, target_name=user.full_name)
    await message.answer(
        f"👤 <b>{user.full_name}</b>\n"
        f"📱 {user.phone}\n"
        f"Поточна роль: {role_label}\n\n"
        "Оберіть нову роль:",
        reply_markup=appoint_role_keyboard(),
        parse_mode="HTML"
    )
    await state.set_state(AppointVice.confirm)

@router.callback_query(F.data.startswith("appoint_role:"))
async def appoint_choose_role(callback: CallbackQuery, state: FSMContext):
    if not is_director(callback.from_user.id):
        await callback.answer("❌ Тільки директор.", show_alert=True)
        return

    role_key = callback.data.split(":")[1]
    data = await state.get_data()
    target_id = data.get("target_id")
    target_name = data.get("target_name")

    if not target_id:
        await callback.answer("Помилка — почніть знову.", show_alert=True)
        await state.clear()
        return

    role_map = {
        "vice": UserRole.VICE_PRINCIPAL,
        "secretary": UserRole.SECRETARY,
    }
    new_role = role_map.get(role_key)
    if not new_role:
        await callback.answer("Невідома роль.", show_alert=True)
        return

    role_labels = {
        UserRole.VICE_PRINCIPAL: "Заступник директора",
        UserRole.SECRETARY: "Секретар",
    }

    await callback.message.edit_text(
        f"Призначити <b>{target_name}</b> як <b>{role_labels[new_role]}</b>?",
        reply_markup=confirm_appoint_keyboard(target_id, role_key),
        parse_mode="HTML"
    )
    await state.update_data(new_role=role_key)
    await callback.answer()

@router.callback_query(F.data.startswith("appoint_vice:") | F.data.startswith("appoint_secretary:"))
async def confirm_appoint(callback: CallbackQuery, state: FSMContext):
    if not is_director(callback.from_user.id):
        await callback.answer("❌ Тільки директор.", show_alert=True)
        return

    parts = callback.data.split(":")
    role_key = parts[0].replace("appoint_", "")
    target_telegram_id = int(parts[1])

    role_map = {
        "vice": UserRole.VICE_PRINCIPAL,
        "secretary": UserRole.SECRETARY,
    }
    new_role = role_map.get(role_key, UserRole.TEACHER)
    role_labels = {
        UserRole.VICE_PRINCIPAL: "Заступник директора 👔",
        UserRole.SECRETARY: "Секретар 👩‍💼",
    }

    async with SessionLocal() as session:
        result = await session.execute(
            select(User).where(User.telegram_id == target_telegram_id)
        )
        target_user = result.scalar_one_or_none()

        director_result = await session.execute(
            select(User).where(User.telegram_id == callback.from_user.id)
        )
        director = director_result.scalar_one_or_none()

        if not target_user:
            await callback.answer("❌ Користувача не знайдено.", show_alert=True)
            return

        target_user.role = new_role
        if director:
            target_user.appointed_by = director.id
        await session.commit()
        name = target_user.full_name

    await callback.message.edit_text(
        f"✅ <b>{name}</b> призначено як <b>{role_labels.get(new_role, new_role.value)}</b>!",
        parse_mode="HTML"
    )
    await callback.answer()
    await state.clear()

    # Повідомити призначеного
    try:
        await callback.bot.send_message(
            target_telegram_id,
            f"🎉 Вас призначено <b>{role_labels.get(new_role, '')}</b> у системі!\n"
            "Натисніть /start щоб оновити меню.",
            parse_mode="HTML"
        )
    except Exception:
        pass

    await callback.message.answer("Меню:", reply_markup=director_main_menu())

@router.callback_query(F.data == "appoint_cancel")
async def cancel_appoint(callback: CallbackQuery, state: FSMContext):
    await callback.message.edit_text("Скасовано.")
    await callback.answer()
    await state.clear()
    await callback.message.answer("Меню:", reply_markup=director_main_menu())

# ─────────────────────────────────────────────────────────
# 2. ВІДПУСТКИ — директор переглядає всі схвалені
# ─────────────────────────────────────────────────────────

@router.message(F.text == "🏖 Відпустки")
async def director_vacations(message: Message):
    if not is_director(message.from_user.id):
        await message.answer("❌ Тільки для директора.")
        return

    async with SessionLocal() as session:
        result = await session.execute(
            select(Vacation)
            .options(selectinload(Vacation.user))
            .where(Vacation.status == VacationStatus.APPROVED)
            .order_by(Vacation.start_date.desc())
        )
        vacations = result.scalars().all()

    if not vacations:
        await message.answer("📭 Схвалених відпусток немає.")
        return

    lines = []
    for i, v in enumerate(vacations, 1):
        lines.append(
            f"{i}. <b>{v.user.full_name}</b>\n"
            f"   📅 {v.start_date} – {v.end_date} ({v.days_count} дн.)"
        )

    await message.answer(
        "🏖 <b>Схвалені відпустки:</b>\n\n" + "\n\n".join(lines),
        parse_mode="HTML"
    )

# ─────────────────────────────────────────────────────────
# 3. СТВОРИТИ НАКАЗИ — список схвалених відпусток → docx
# ─────────────────────────────────────────────────────────

@router.message(F.text == "📋 Створити накази")
async def director_orders_menu(message: Message):
    if not is_director(message.from_user.id):
        await message.answer("❌ Тільки для директора.")
        return

    async with SessionLocal() as session:
        result = await session.execute(
            select(Vacation)
            .options(selectinload(Vacation.user))
            .where(Vacation.status == VacationStatus.APPROVED)
            .order_by(Vacation.start_date.desc())
        )
        vacations = result.scalars().all()

    if not vacations:
        await message.answer("📭 Немає схвалених відпусток для генерації наказів.")
        return

    await message.answer(
        "📋 <b>Оберіть відпустку для якої згенерувати наказ:</b>\n"
        "(натисніть на ім'я у списку нижче)",
        reply_markup=approved_vacations_keyboard(vacations),
        parse_mode="HTML"
    )

@router.callback_query(F.data.startswith("gen_order:"))
async def generate_vacation_order(callback: CallbackQuery):
    if not is_director(callback.from_user.id):
        await callback.answer("❌ Тільки директор.", show_alert=True)
        return

    vacation_id = int(callback.data.split(":")[1])

    async with SessionLocal() as session:
        result = await session.execute(
            select(Vacation)
            .options(selectinload(Vacation.user))
            .where(Vacation.id == vacation_id)
        )
        vacation = result.scalar_one_or_none()

        director_result = await session.execute(
            select(User).where(User.telegram_id == callback.from_user.id)
        )
        director = director_result.scalar_one_or_none()

    if not vacation:
        await callback.answer("❌ Відпустку не знайдено.", show_alert=True)
        return

    await callback.answer("⏳ Генерую наказ...")

    try:
        file_path = await _generate_vacation_order_docx(vacation, director)
        doc_file = FSInputFile(file_path)
        await callback.message.answer_document(
            doc_file,
            caption=(
                f"📄 Наказ про відпустку\n"
                f"👤 {vacation.user.full_name}\n"
                f"📅 {vacation.start_date} – {vacation.end_date} ({vacation.days_count} дн.)"
            )
        )
    except Exception as e:
        await callback.message.answer(f"❌ Помилка генерації: {e}")

# ─────────────────────────────────────────────────────────
# Генерація docx наказу на відпустку
# ─────────────────────────────────────────────────────────

async def _generate_vacation_order_docx(vacation, director) -> str:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    employee = vacation.user
    director_name = director.full_name if director else "Директор"

    # Форматуємо дати
    months_uk = {
        1: "січня", 2: "лютого", 3: "березня", 4: "квітня",
        5: "травня", 6: "червня", 7: "липня", 8: "серпня",
        9: "вересня", 10: "жовтня", 11: "листопада", 12: "грудня"
    }
    start = vacation.start_date
    end = vacation.end_date
    today = datetime.now()

    start_str = f"{start.day} {months_uk[start.month]} {start.year} року"
    end_str = f"{end.day} {months_uk[end.month]} {end.year} року"
    today_str = f"{today.day} {months_uk[today.month]} {today.year}"

    # Визначаємо навчальний рік
    if today.month >= 9:
        school_year = f"{today.year}-{today.year+1}"
    else:
        school_year = f"{today.year-1}-{today.year}"

    doc = Document()

    # Поля
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    def add_paragraph(text="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        return p

    # Заголовок школи
    add_paragraph(
        "Павлоградський міський ліцей",
        bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14
    )

    # Номер та дата наказу
    order_num_p = doc.add_paragraph()
    order_num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = order_num_p.add_run(f"НАКАЗ № ___\n{today_str} р.")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # Тема наказу
    add_paragraph(
        f"Про надання щорічної відпустки\n{employee.full_name}",
        bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12
    )

    doc.add_paragraph()

    # Підстава
    basis_p = doc.add_paragraph()
    basis_run = basis_p.add_run(
        f"Відповідно до статті 74 Кодексу законів про працю України, "
        f"статті 4 Закону України «Про відпустки», "
        f"на підставі заяви {employee.full_name},"
    )
    basis_run.font.size = Pt(12)
    basis_run.font.name = "Times New Roman"

    doc.add_paragraph()

    # НАКАЗУЮ
    add_paragraph("НАКАЗУЮ:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)

    doc.add_paragraph()

    # Основна частина
    main_p = doc.add_paragraph()
    main_run = main_p.add_run(
        f"1. Надати {employee.full_name} щорічну основну оплачувану відпустку "
        f"тривалістю {vacation.days_count} ({"".join(_num_to_words(vacation.days_count))}) календарних дні(в) "
        f"з {start_str} по {end_str}."
    )
    main_run.font.size = Pt(12)
    main_run.font.name = "Times New Roman"

    doc.add_paragraph()

    control_p = doc.add_paragraph()
    control_run = control_p.add_run(
        "2. Контроль за виконанням даного наказу залишаю за собою."
    )
    control_run.font.size = Pt(12)
    control_run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph()

    # Підпис директора
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Прізвище великими літерами
    parts = director_name.strip().split()
    if len(parts) >= 2:
        surname_upper = parts[-1].upper()
        first_parts = " ".join(parts[:-1])
        sign_text = f"Директор{' ' * 40}{first_parts} {surname_upper}"
    else:
        sign_text = f"Директор{' ' * 40}{director_name.upper()}"
    sign_run = sign_p.add_run(sign_text)
    sign_run.font.size = Pt(12)
    sign_run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph()

    # Ознайомлений
    add_paragraph("З наказом ознайомлений(а):", size=12)
    ack_p = doc.add_paragraph()
    parts_emp = employee.full_name.strip().split()
    if len(parts_emp) >= 2:
        surname_upper = parts_emp[-1].upper() if len(parts_emp) == 1 else parts_emp[0].upper()
        # ПІБ: Прізвище ІВАНОВ, Ім'я По батькові
        ack_name = f"{parts_emp[0].upper()} {' '.join(parts_emp[1:])}"
    else:
        ack_name = employee.full_name.upper()
    ack_run = ack_p.add_run(f"{ack_name}{' ' * 20}_______________  /{today_str}/")
    ack_run.font.size = Pt(12)
    ack_run.font.name = "Times New Roman"

    # Зберігаємо
    os.makedirs("/tmp/orders", exist_ok=True)
    safe_name = re.sub(r"[^\w]", "_", employee.full_name)
    file_path = f"/tmp/orders/nakaz_vidpustka_{safe_name}_{start.year}_{start.month:02d}.docx"
    doc.save(file_path)
    return file_path


def _num_to_words(n: int) -> str:
    """Прості числа словами (для кількості днів)"""
    words = {
        1: "один", 2: "два", 3: "три", 4: "чотири", 5: "п'ять",
        6: "шість", 7: "сім", 8: "вісім", 9: "дев'ять", 10: "десять",
        11: "одинадцять", 12: "дванадцять", 13: "тринадцять", 14: "чотирнадцять",
        15: "п'ятнадцять", 16: "шістнадцять", 17: "сімнадцять", 18: "вісімнадцять",
        19: "дев'ятнадцять", 20: "двадцять", 21: "двадцять один", 22: "двадцять два",
        23: "двадцять три", 24: "двадцять чотири", 25: "двадцять п'ять",
        26: "двадцять шість", 27: "двадцять сім", 28: "двадцять вісім",
        29: "двадцять дев'ять", 30: "тридцять",
    }
    return words.get(n, str(n))
